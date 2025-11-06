import base64
from configparser import ConfigParser
import openai
import os
from spire.doc import *
from spire.doc.common import *
import re


def get_config_string(section: str, key: str = None):
    """
    取得config.ini裡面的設定值
    """
    config = ConfigParser()
    config.read('config.ini')
    if key is None:
        return config[section]
    else:
        return config[section][key]

def call_openai_api(messages, model_name="gpt-4o"):
    '''
    呼叫 OpenAI API
    :param api_key: OpenAI API 金鑰
    :param messages: 要傳遞給 API 的訊息
    :return: API 回傳的結果
    '''
    # 呼叫 OpenAI API
    # 設定 OpenAI API 金鑰
    openai.api_key = get_config_string('AzureOpenAI','api_key')
    openai.azure_endpoint = get_config_string('AzureOpenAI','api_endpoint')
    openai.api_type = get_config_string('AzureOpenAI','api_type')
    openai.api_version = get_config_string('AzureOpenAI','api_version')
    response = openai.chat.completions.create(
        model=model_name, messages=messages)
    return response

def img_to_base64string(img_path):
    '''
    將圖片轉換成base64格式
    :param img_path: 圖片的檔案路徑
    :return: base64格式的字串
    '''
    try:
        with open(img_path, "rb") as img_file:
            base64_string = base64.b64encode(img_file.read()).decode('utf-8')
        return base64_string
    except FileNotFoundError:
        print(f"檔案未找到: {img_path}")
        return None
    except Exception as e:
        print(f"發生錯誤: {e}")
        return None

def extract_text_from_image(img_path):
    '''
    使用 OpenAI 的 gpt-4o 模型擷取圖片中的文字並轉換為 Markdown 格式
    :param img_path: 圖片的檔案路徑
    :return: Markdown 格式的文字
    '''
    # 將圖片轉換為 base64 字串
    base64_string = img_to_base64string(img_path)
    if not base64_string:
        return None

    try:

        messages = [
            {
                "role": "user",
                "content": [
                    { 
                        "type": "text", 
                        "text": """
                                #zh-tw
                                請幫我擷取圖片中的文字，然後轉成Markdown格式
                                圖片中有特別標示顏色即為正確答案，這部分請用粗體來表示。格式請參考範例
                                請用Markdown格式來呈現，並且不要有任何其他的文字或說明
                                [正確答案格式範例-開始]
                                ### Q1. 生成式AI的主要特點為何？
                                1. 生成式AI無法生成新的內容，只能回答特定問題  
                                2. **生成式AI可以理解和生成自然語言，並能創造出新的內容**  
                                3. 生成式AI只能在特定領域中工作，無法跨領域生成內容  
                                4. 生成式AI無法理解影像或視覺內容  
                                [正確答案格式範例-結束]
                                """
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:image/png;base64,{base64_string}",
                        },
                    },
                ],
            }
        ]


        response = call_openai_api(messages)
        extracted_text = response.choices[0].message.content
        print(f"總token數: {response.usage.total_tokens}")

        # 將文字轉換為 Markdown 格式
        markdown_text = f"```\n{extracted_text}\n```"
        return markdown_text
    except Exception as e:
        print(f"發生錯誤: {e}")
        return None

def get_paths_from_folder(folder_path, 
                        supported_extensions=('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff')):
    '''
    獲取資料夾中所有有支援檔案的路徑
    :param folder_path: 資料夾的路徑
    :return: 包含所有有支援檔案路徑的列表
    '''
    # supported_extensions = ('.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.md')
    image_paths = []

    try:
        for root, _, files in os.walk(folder_path):
            for file in files:
                if file.lower().endswith(supported_extensions):
                    image_paths.append(os.path.join(root, file))
        return image_paths
    except Exception as e:
        print(f"發生錯誤: {e}")
        return []

def ocr_with_llm():
    '''
    主函數，負責擷取資料夾中的圖片並轉換為 Markdown 格式
    '''
    img_paths = get_paths_from_folder("img")
    print(f"找到 {len(img_paths)} 張圖片")
    print("開始擷取文字...")
    for img_path in img_paths:
        markdown_text = extract_text_from_image(img_path)
        markdown_text = re.sub(r'```|markdown', '', markdown_text)
        # print(f"擷取到的文字:\n {markdown_text.strip('\n').strip()}")
        if markdown_text:
            # 儲存擷取到的文字
            output_path = os.path.splitext(img_path)[0] + ".md"
            output_path = output_path.replace("img", "txt")
            # 確保資料夾存在
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            # 儲存 Markdown 格式的文字
            with open(output_path, "w", encoding="utf-8") as output_file:
                output_file.write(markdown_text)
            print(f"擷取成功: {output_path}")
        else:
            print("無法擷取文字")
    print("擷取完成")

def markdown_to_docx(root_path):
    '''
    將 Markdown 格式的文字轉換為 Docx 格式
    :param markdown_text: Markdown 格式的文字
    :return: Docx 格式的文件
    '''
    md_filePath = get_paths_from_folder(root_path, ('.md'))
    print(f"找到 {len(md_filePath)} 個 Markdown 檔案")
    print("開始轉換為 Docx 格式...")
    
    # Create a Document object
    document = Document()
    for md_path in md_filePath:
        print(f"正在轉換: {md_path}")
        # Load the Markdown file
        document.LoadFromFile(md_path)

        output_path = os.path.splitext(md_path)[0] + ".docx"
        output_path = output_path.replace("txt", "docx")
        # 確保資料夾存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        # Save it as a docx file
        document.SaveToFile(output_path, FileFormat.Docx2019)
    # Dispose resources
    document.Dispose()
    print("轉換完成")

def docx_postprocess(docx_path):
    '''
    對 Docx 格式的文件進行後處理
    :param docx_path: Docx 檔案的路徑
    :return: None
    '''
    from docx import Document

    docx_paths = get_paths_from_folder(docx_path, ('.docx'))
    print(f"找到 {len(docx_paths)} 個 Docx 檔案")
    print("開始後處理...")
    for docx_path in docx_paths:
        print(f"正在後處理: {docx_path}")
        # Load the Docx file
        document = Document(docx_path)

        # 刪除第一行的文字
        if len(document.paragraphs) > 0:
            document.paragraphs[0].clear()
        # Save it as a docx file
        document.save(docx_path)
    print("後處理完成")

def merge_all_docx_files(docx_path_root):
    '''
    將所有 Docx 格式的文件合併為一個文件
    :param docx_path: Docx 檔案的路徑
    :return: None
    '''
    from docx import Document

    
    docx_paths = get_paths_from_folder(docx_path_root, ('.docx'))
    print(f"找到 {len(docx_paths)} 個 Docx 檔案")
    print("開始合併...")

    # Create a new Document object
    merged_document = Document()
    dir_list = os.listdir(docx_path_root)
    
    for dir in dir_list:
        path = get_paths_from_folder(f"{docx_path_root}/{dir}", ('.docx'))
        print(f"{dir}資料夾中找到 {len(path)} 個 Docx 檔案")
        merged_document.add_heading(dir, level=2)
        for docx_path in path:
            print(f"正在合併: {docx_path}")
            # Load the Docx file
            document = Document(docx_path)
            # Append the content to the merged document
            for paragraph in document.paragraphs:            
                merged_document.add_paragraph(text=paragraph.text, style=paragraph.style)
            # Append a page break after each document
            merged_document.add_page_break()
    
    # Save the merged document
    merged_document.save("生成式AI能力課程講師手冊2025(18週)-題目.docx")
    print("合併完成") 

def markdown_to_zuvioIRS_quizbank_csv(markdown_text):
    '''
    將 Markdown 格式的文字轉換為 Zuvio IRS QuizBank CSV 格式
    :param markdown_text: Markdown 格式的文字
    :return: Zuvio IRS QuizBank CSV 格式的文件
    '''
    messages = [
        {
            "role": "user",
            "content": f"""
                #zh-tw
                以下是一個Markdown格式的內容，markdown格式內容中，選項有標示粗體即為正確答案，請根據標的格式的範例以及注意事項以及注意事項來做轉換，不要產生任何東西
                
                [Markdown格式的內容開始]
                {markdown_text}
                [Markdown格式的內容結束]

                [標的格式的範例開始]
                2,Q1. AI文字生成原理是基於什麼技術？,,2,1. 機器學習,2. 自然語言處理,3. 自回歸分析,4. 深度學習
                [標的格式的範例結束]

                [注意事項 開始]
                1. 第一欄為固定值為2
                2. 第二欄為題目
                3. 第三欄為空值
                4. 第四欄為正確答案
                5. 選項請按順序排列，並且用逗號隔開
                [注意事項 結束]
            """
        }
    ]
    response = call_openai_api(messages)
    extracted_text = response.choices[0].message.content
    print(f"總token數: {response.usage.total_tokens}")
    # 將文字轉換為 CSV 格式字串並存檔
    csv_text = f"{extracted_text}"
    # 儲存 CSV 格式的文字
    output_path = "output.csv"
    with open(output_path, "a", encoding="utf-8") as output_file:
        output_file.write(csv_text)
    print(f"儲存成功: {output_path}")
    
def process_markdown_files(root_path):
    '''
    主函數，負責處理資料夾中的 Markdown 檔案
    :param root_path: 資料夾的路徑
    :return: None
    '''
    md_filePath = get_paths_from_folder(root_path, ('.md'))
    print(f"找到 {len(md_filePath)} 個 Markdown 檔案")
    print("開始轉換為 Zuvio IRS QuizBank CSV 格式...")
    
    for md_path in md_filePath:
        print(f"正在轉換: {md_path}")
        # Load the Markdown file
        with open(md_path, "r", encoding="utf-8") as file:
            markdown_text = file.read()
        # Convert to Zuvio IRS QuizBank CSV format
        markdown_to_zuvioIRS_quizbank_csv(markdown_text)

def markdown_to_moodle_quiz_aiken_format(markdown_text, txt_path=None):
    '''
    將 Markdown 格式的文字轉換為 Moodle Quiz Aiken 格式
    :param markdown_text: Markdown 格式的文字
    :param txt_path: 輸出檔案的路徑
    :return: Moodle Quiz Aiken 格式的文件
    '''
    messages = [
        {
            "role": "user",
            "content": f"""
                以下是一段以 Markdown 格式撰寫的內容。
                其中每題的選項中，以粗體標示者為正確答案。
                請根據下方「範例格式」與「注意事項」，將該內容轉換為指定格式。
                不要產生任何額外文字、說明或翻譯。
                
                [Markdown格式的內容開始]
                
                {markdown_text}
                
                [Markdown格式的內容結束]

                [標的格式的範例開始]
                
                AI文字生成原理是基於什麼技術？
                A. 機器學習  
                B. 自然語言處理
                C. 自回歸分析  
                D. 深度學習
                ANSWER: B

                下列何者是生成式AI的特色？
                A. 無法客製化  
                B. 多領域應用  
                C. 不斷退步  
                D. 無法提供創意和靈感
                ANSWER: B
                
                [標的格式的範例結束]

                [注意事項開始]
                
                1. 題幹前不加題號（匯入 Moodle 後系統會自動生成）。各題之間請插入一空行。
                2. 題幹內容必須置於同一行，不可使用換行標籤。
                3. 每個選項須以大寫英文字母 A、B、C、D 為標號，後接一個「.」或「)」，再空一格後接選項內容。
                4. 正確答案須放在最後一個選項下方，以「ANSWER: 」開頭（半形冒號），後接代表答案的大寫字母。
                5. 不要翻譯題目內容。
                6. 不要輸出任何額外文字或解釋。
                
                [注意事項結束]
            """
        }
    ]
    response = call_openai_api(messages)
    extracted_text = response.choices[0].message.content
    print(f"總token數: {response.usage.total_tokens}")
    # 將文字轉換為 Moodle Quiz 格式字串並存檔
    moodle_quiz_text = f"{extracted_text}"
    
    if txt_path is None:
        output_path = "output.txt"
    else:
        output_path = os.path.dirname(txt_path) + ".txt"
        output_path = output_path.replace("txt", "moodle_quiz_aiken")
        os.makedirs(os.path.dirname(output_path), exist_ok=True) # 確保資料夾存在

    # 儲存 Moodle Quiz 格式的文字
    with open(output_path, "a", encoding="utf-8") as output_file:
        output_file.write(moodle_quiz_text)
    print(f"儲存成功: {output_path}")

def process_markdown_files_to_moodle_quiz_format(root_path):
    '''
    主函數，負責處理資料夾中的 Markdown 檔案並轉換為 Moodle Quiz 格式
    :param root_path: 資料夾的路徑
    :return: None
    '''
    md_filePath = get_paths_from_folder(root_path, ('.md'))
    print(f"找到 {len(md_filePath)} 個 Markdown 檔案")
    print("開始轉換為 Moodle Quiz 格式...")
    
    for md_path in md_filePath:
        print(f"正在轉換: {md_path}")
        # Load the Markdown file
        with open(md_path, "r", encoding="utf-8") as file:
            markdown_text = file.read()
        # Convert to Moodle Quiz format
        markdown_to_moodle_quiz_aiken_format(markdown_text, md_path)

if __name__ == '__main__':
    # ocr_with_llm()
    process_markdown_files_to_moodle_quiz_format('txt')
    # markdown_to_docx('txt')
    # docx_postprocess('docx')
    # merge_all_docx_files('docx')
    # process_markdown_files('txt')